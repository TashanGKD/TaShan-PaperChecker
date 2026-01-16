from .base_extractor import BaseExtractor
from models.document import Document, Citation, Reference
import docx
import re

class WordExtractor(BaseExtractor):
    """Word文档提取器"""
    
    def validate_file(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc'))
    
    def extract(self, file_path: str) -> Document:
        """提取Word文档内容"""
        doc = docx.Document(file_path)
        
        # 提取正文内容
        paragraphs = []
        for paragraph in doc.paragraphs:
            paragraphs.append(paragraph.text)
        
        # 提取表格内容
        tables_content = []
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    tables_content.append(cell.text)
        
        # 提取引文
        citations = self._extract_citations(paragraphs, file_path)
        
        # 提取参考文献
        references = self._extract_references(paragraphs)
        
        # 构建文档对象
        return Document(
            content=paragraphs,
            tables=tables_content,
            citations=citations,
            references=references,
            metadata={'file_type': 'docx', 'file_path': file_path}
        )
    
    def _extract_citations(self, paragraphs: list, file_path: str) -> list:
        """提取引文"""
        all_citations = []
        processed_citations = set()  # 用于去重
        
        # 从extractor模块导入AI增强的引用提取功能
        try:
            from .ai_extractor import extract_references, extract_western_references
            AI_ENHANCED_EXTRACTION_AVAILABLE = True
        except ImportError:
            AI_ENHANCED_EXTRACTION_AVAILABLE = False
        
        # 如果AI增强提取功能可用，使用AI提取作者年份格式的引用
        if AI_ENHANCED_EXTRACTION_AVAILABLE:
            try:
                # 准备AI优化的配置（这里可以使用从配置文件传入的配置参数）
                ai_config = {
                    "api_key": "your-api-key",  # 可以从配置文件传入
                    "model_name": "qwen-plus"   # 可以从配置文件传入
                }
                
                # 提取作者年份格式的引用
                author_year_citations = extract_references(file_path, ai_config)
                western_citations = extract_western_references(file_path, ai_config)
                
                # 将AI提取的引用添加到列表中
                for citation_text in author_year_citations + western_citations:
                    if citation_text not in processed_citations:
                        # 创建引用对象，格式化为AI提取的格式
                        citation_obj = Citation(
                            text=citation_text,
                            format_type='author_year',
                            context="AI提取"
                        )
                        all_citations.append(citation_obj)
                        processed_citations.add(citation_text)
            except Exception as e:
                print(f"AI增强引用提取失败: {e}")
        
        # 提取数字格式引用 [1], [2-5] 等（作为补充）
        for paragraph in paragraphs:
            # 提取数字格式引用 [1], [2-5] 等
            number_citations = re.findall(r'\[\d+(?:-\d+)?\]', paragraph)
            for citation in number_citations:
                if citation not in processed_citations:
                    all_citations.append(Citation(text=citation, format_type='number', context=paragraph))
                    processed_citations.add(citation)
            
            # 提取作者年份格式引用作为补充（以防AI提取遗漏）
            author_year_citations = self._extract_author_year_citations(paragraph)
            for citation in author_year_citations:
                if citation.text not in processed_citations:
                    all_citations.append(citation)
                    processed_citations.add(citation.text)
        
        return all_citations
    
    def _extract_author_year_citations(self, paragraph: str) -> list:
        """提取作者年份格式引用"""
        citations = []
        # 中文格式：张三（2024）
        chinese_pattern = r'([\u4e00-\u9fa5\w\s]+)（(\d{4})）'
        chinese_matches = re.findall(chinese_pattern, paragraph)
        for author, year in chinese_matches:
            citations.append(Citation(text=f"{author}（{year}）", 
                                    format_type='author_year', 
                                    author=author.strip(), 
                                    year=year, 
                                    context=paragraph))
        
        # 英文格式：Smith (2020) 或 Smith（2020）
        english_pattern = r'([A-Za-z\s\.\-&]+)\s*[（\(](\d{4})[）\)]'
        english_matches = re.findall(english_pattern, paragraph)
        for author, year in english_matches:
            citations.append(Citation(text=f"{author.strip()} ({year})", 
                                    format_type='author_year', 
                                    author=author.strip(), 
                                    year=year, 
                                    context=paragraph))
        
        return citations
    
    def _extract_references(self, paragraphs: list) -> list:
        """提取参考文献"""
        references = []
        references_start = -1
        
        # 查找参考文献部分
        for i, paragraph in enumerate(paragraphs):
            if any(keyword in paragraph for keyword in ['参考文献', 'References', 'REFERENCES']):
                references_start = i
                break
        
        if references_start != -1:
            # 提取参考文献条目（从标题之后开始）
            for i in range(references_start + 1, len(paragraphs)):
                text = paragraphs[i].strip()
                if text:
                    # 检查是否为参考文献条目
                    if self._is_reference_entry(text):
                        # 提取DOI、URL等信息
                        doi = self._extract_doi(text)
                        url = self._extract_url(text)
                        
                        # 从参考文献文本中提取作者和年份
                        from core.checker.citation_checking.reference_mapper import extract_author_year_from_reference
                        extracted_author, extracted_year = extract_author_year_from_reference(text)
                        
                        references.append(Reference(
                            text=text,
                            author=extracted_author,
                            year=extracted_year
                        ))
                    # 如果遇到结束标记则停止
                    elif self._is_end_marker(text):
                        break
        
        return references
    
    def _is_reference_entry(self, text: str) -> bool:
        """判断是否为参考文献条目"""
        # 检查是否包含年份和基本结构
        has_year = bool(re.search(r'\d{4}', text))
        has_basic_structure = len(text) > 10 and ('.' in text or '[' in text)
        
        # 排除非参考文献的条目
        exclude_keywords = ['附录', '致谢', '作者简历', '图', '表', '目录']
        has_exclude_keyword = any(keyword in text for keyword in exclude_keywords)
        
        return has_year and has_basic_structure and not has_exclude_keyword
    
    def _is_end_marker(self, text: str) -> bool:
        """判断是否为参考文献部分结束标记"""
        end_keywords = ['附录', '致谢', '作者简历', 'Appendix', 'Acknowledgements']
        return any(keyword in text for keyword in end_keywords) and len(text) < 20
    
    def _extract_doi(self, text: str) -> str:
        """提取DOI"""
        doi_pattern = r'doi:\s*([^\s,.;\)]+)|DOI:\s*([^\s,.;\)]+)'
        match = re.search(doi_pattern, text, re.IGNORECASE)
        if match:
            return match.group(1) or match.group(2)
        return None
    
    def _extract_url(self, text: str) -> str:
        """提取URL"""
        url_pattern = r'https?://[^\s,.;\)]+'
        match = re.search(url_pattern, text)
        return match.group(0) if match else None